import os
import json
import zipfile
from docx import Document
import re

def extract_images(docx_path, output_dir, prefix):
    images = []
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        for file in zip_ref.namelist():
            if file.startswith('word/media/') and not file.endswith('/'):
                extension = file.split('.')[-1]
                image_name = f"{prefix}_{len(images) + 1}.{extension}"
                image_path = os.path.join(output_dir, image_name)
                with open(image_path, 'wb') as f:
                    f.write(zip_ref.read(file))
                images.append(f"/products/media/{image_name}")
    return images

def process_file(docx_path, category_prefix="03"):
    doc = Document(docx_path)
    filename = os.path.basename(docx_path)
    filename_base = os.path.splitext(filename)[0]
    
    # 1. Product Name
    product_name = filename_base
    if doc.paragraphs and len(doc.paragraphs[0].text) < 50:
        title_text = doc.paragraphs[0].text.strip()
        if title_text and title_text not in ["产品简介", "产品介绍", "产品参数", "技术参数", "参数清单"]:
            product_name = title_text
    
    # Rest of the parsing logic...
    # (Leaving it the same as existing replace_file_content already improved it)

    # 2. Parameters (Do this first to help with summary/detail)
    parameters = {}
    
    # Extract from tables
    for table in doc.tables:
        if len(table.rows) >= 2:
            # Check if it's a header-value table
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            # Handle multi-line headers or sub-headers
            if len(table.rows) >= 3 and table.rows[1].cells[0].text.strip() == table.rows[0].cells[0].text.strip():
                # Might have merged cells or sub-headers in row 1
                for i, cell in enumerate(table.rows[1].cells):
                    txt = cell.text.strip()
                    if txt and txt not in headers:
                        headers[i] = txt
            
            for row_idx in range(1, len(table.rows)):
                row_cells = [cell.text.strip() for cell in table.rows[row_idx].cells]
                # If row is just repeating headers, skip
                if any(c in headers for c in row_cells if c):
                    if row_cells[0] != headers[0]: # Unless first cell is different (like a model name)
                        pass
                    else:
                        continue
                
                for i, val in enumerate(row_cells):
                    if i < len(headers):
                        key = headers[i]
                        if key and val and key != val:
                            # If key already exists, maybe it's a new row for same keys
                            if key in parameters and parameters[key] != val:
                                parameters[f"{key}_{row_idx}"] = val
                            else:
                                parameters[key] = val
                            
    # Extract from paragraphs (Key: Value)
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt: continue
        # Match both ： and :
        match = re.match(r'^([^：:]+)[：:](.+)$', txt)
        if match:
            k, v = match.groups()
            k = k.strip()
            v = v.strip()
            if len(k) < 30 and k not in ["http", "https"]:
                parameters[k] = v

    # 3. Summary
    # Extract all text to find introduction
    all_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    summary_text = ""
    found_intro = False
    for i, line in enumerate(all_lines):
        if "产品介绍" in line or "Product Introduction" in line or "简介" in line:
            # Start collecting from next line
            collected = []
            char_count = 0
            for next_line in all_lines[i+1:]:
                # Stop if we hit a new section or heavy parameters
                if any(x in next_line for x in ["技术指标", "规格参数", "功能特点"]):
                    break
                # Filter out pure English lines if they are just translations
                collected.append(next_line)
                char_count += len(next_line)
                if char_count > 300: break
            summary_text = " ".join(collected)
            found_intro = True
            break
            
    if not found_intro:
        # Just grab the first few non-empty paragraphs that aren't titles
        summary_text = " ".join(all_lines[1:4]) if len(all_lines) > 1 else product_name

    # Clean summary: remove extra spaces, limit to ~150 chars
    summary_text = re.sub(r'\s+', ' ', summary_text).strip()
    if len(summary_text) > 200:
        summary_text = summary_text[:197] + "..."
    elif len(summary_text) < 50:
        # If too short, keep original or expand
        pass
    summary = summary_text

    # 4. Detail HTML
    detail_lines = []
    media_dir = "/Users/mattchyi/Documents/Project/fc/public/products/media"
    image_prefix = f"{category_prefix}_{filename_base}"
    image_paths = extract_images(docx_path, media_dir, image_prefix)
    
    current_list = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt: continue
        
        # Heading detection
        is_heading = False
        if len(txt) < 30 and (txt.endswith("指标") or txt.endswith("参数") or txt.endswith("特点") or "介绍" in txt or "规格" in txt):
            is_heading = True
        
        if is_heading:
            if current_list:
                detail_lines.append("<ul><li>" + "</li><li>".join(current_list) + "</li></ul>")
                current_list = []
            detail_lines.append(f"<h4>{txt}</h4>")
        else:
            current_list.append(txt)

    if current_list:
        detail_lines.append("<ul><li>" + "</li><li>".join(current_list) + "</li></ul>")
    
    if image_paths:
        img_html = '<div class="product-images">'
        for img in image_paths:
            img_html += f'<img src="{img}" alt="{product_name}" style="max-width:100%; margin-bottom:10px; display:block;" />'
        img_html += '</div>'
        detail_lines.append(img_html)
        
    detail_html = "".join(detail_lines)
    
    result = {
        "product_name": product_name,
        "summary": summary,
        "parameters": parameters,
        "detail_html": detail_html
    }
    
    output_path = docx_path.replace(".docx", ".json")
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    
    return result

if __name__ == "__main__":
    files = [
        "/Users/mattchyi/Documents/Project/fc/网站资料/04工程补给/贝雷钢桥/01工程补给产品贝雷桥.docx",
        "/Users/mattchyi/Documents/Project/fc/网站资料/05野战医院/01箱组式医疗救治系统.docx",
        "/Users/mattchyi/Documents/Project/fc/网站资料/05野战医院/02智能化可移动式多功能方舱医院.docx",
        "/Users/mattchyi/Documents/Project/fc/网站资料/06要地防护/FC-DMA型多波段光电转台.docx",
        "/Users/mattchyi/Documents/Project/fc/网站资料/06要地防护/FC-DMS10系列智能电子哨兵.docx",
        "/Users/mattchyi/Documents/Project/fc/网站资料/06要地防护/FC-DTVC系列双波段热成像高速球.docx",
        "/Users/mattchyi/Documents/Project/fc/网站资料/06要地防护/FC-RC系列高清激光摄像机.docx",
        "/Users/mattchyi/Documents/Project/fc/网站资料/06要地防护/FC-RDS500-4R型雷视融合系统.docx",
        "/Users/mattchyi/Documents/Project/fc/网站资料/06要地防护/FC-TTVC系列智能多波段摄像机.docx"
    ]
    for f in files:
        if os.path.exists(f):
            # Extract category prefix from path (e.g., "04" from "04工程补给")
            category_prefix = "01"
            parts = f.split("/")
            for part in parts:
                if len(part) >= 2 and part[:2].isdigit():
                    category_prefix = part[:2]
                    break
            process_file(f, category_prefix=category_prefix)
        else:
            print(f"File not found: {f}")

