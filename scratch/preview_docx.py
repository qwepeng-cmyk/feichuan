import sys
from docx import Document

def preview_docx(filepath):
    doc = Document(filepath)
    print(f"--- Previewing: {filepath} ---")
    for para in doc.paragraphs:
        if para.text.strip():
            print(f"P: {para.text[:100]}...")
    for table in doc.tables:
        print(f"--- Table Found ---")
        for row in table.rows:
            print(" | ".join(cell.text.strip().replace('\n', ' ') for cell in row.cells))

preview_docx(sys.argv[1])
