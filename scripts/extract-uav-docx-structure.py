from __future__ import annotations

import json
import sys
from pathlib import Path
from typing import Any

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "网站资料" / "01大无人机"


DOCX_MAP: dict[str, Path] = {
    "multi-rotor-3kg-payload-uav": DATA_DIR / "01多旋翼无人机" / "载重3kg无人机" / "载重3kg无人机.docx",
    "multi-rotor-8kg-payload-uav": DATA_DIR / "01多旋翼无人机" / "载重8kg无人机" / "载重8kg无人机.docx",
    "multi-rotor-20kg-payload-uav": DATA_DIR / "01多旋翼无人机" / "载重20kg无人机" / "载重20kg无人机.docx",
    "multi-rotor-50kg-payload-uav": DATA_DIR / "01多旋翼无人机" / "载重50kg无人机" / "载重50kg无人机.docx",
    "vtol-14kg-mtow-uav": DATA_DIR / "02复合翼无人机" / "起飞重量14kg垂起" / "起飞重量14kg垂起.docx",
    "vtol-26kg-mtow-uav": DATA_DIR / "02复合翼无人机" / "起飞重量26kg垂起" / "垂起26kg无人机.docx",
    "vtol-40kg-mtow-uav": DATA_DIR / "02复合翼无人机" / "起飞重量40kg垂起" / "起飞重量40kg垂起.docx",
    "vtol-64kg-mtow-uav": DATA_DIR / "02复合翼无人机" / "起飞重量64kg垂起" / "起飞重量64kg垂起.docx",
    "vtol-135kg-mtow-uav": DATA_DIR / "02复合翼无人机" / "起飞重量135kg垂起" / "起飞重量135kg垂起.docx",
    "medium-long-range-uav-inspection-system": DATA_DIR / "中远距离无人机巡检.docx",
    "emergency-search-rescue-drone": DATA_DIR / "05新整理无人机2" / "应急搜救无人机.docx",
    "fc-yjtx-01-emergency-communication-drone": DATA_DIR / "05新整理无人机2" / "应急通讯无人机.docx",
    "smart-substation-autonomous-inspection-system": DATA_DIR / "05新整理无人机2" / "智慧变电站无人机全自动巡检系统.docx",
    "fc-sljc-01-water-conservancy-monitoring-drone": DATA_DIR / "05新整理无人机2" / "水利监测无人机.docx",
    "power-tower-inspection-drone": DATA_DIR / "05新整理无人机2" / "电塔巡检无人机.docx",
    "fc-yjzm-01-emergency-lighting-drone": DATA_DIR / "05新整理无人机2" / "系留照明无人机.docx",
    "fc-yjxf-01-aerial-firefighting-drone": DATA_DIR / "05新整理无人机2" / "高层消防无人机.docx",
}


def clean(text: str) -> str:
    return " ".join(text.replace("\u3000", " ").split())


def extract_docx(path: Path) -> dict[str, Any]:
    doc = Document(path)
    paragraphs = []
    for para in doc.paragraphs:
        text = clean(para.text)
        if not text:
            continue
        paragraphs.append(
            {
                "text": text,
                "style": para.style.name if para.style else "",
            }
        )

    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [clean(cell.text) for cell in row.cells]
            while cells and not cells[-1]:
                cells.pop()
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)

    return {
        "source_docx": str(path.relative_to(ROOT)).replace("\\", "/"),
        "paragraphs": paragraphs,
        "tables": tables,
    }


def main() -> None:
    selected = sys.argv[1:]
    handles = selected or list(DOCX_MAP)
    result = {}
    for handle in handles:
        path = DOCX_MAP[handle]
        if not path.exists():
            raise FileNotFoundError(path)
        result[handle] = extract_docx(path)
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
